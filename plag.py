from docx import Document
from google.appengine.api import search

if __name__ == "__main__":
    if sys.args[0] == 0:
        print("Must specify file!")
        return
    #open the docx (and docx only)
    document = Document(sys.args[0])
    #for each paragraph on the docx
    for parag in document.paragraphs:
        #extract the string
        text = parag.text
        #split at whitespace
        splitted = text.split()
        #send to google every 5~10 words and save the url
        #of the first Y results (parallelism preferrable, bandwidth is not a big problem,
        #the old http protocol is)

    # note: check https://cloud.google.com/appengine/training/fts_intro/lesson2
    # for instructions
    # in this case, it will create an index document and fill it with the
    # splitted text

    index = search.Index(name='document')
    fields = [
            search.TextField(name=docs.Document.TEXT, value=text)
    ]

    d = search.Document(fields=fields)
    # need to go back to work, will leave unfinished



    #count the ocurrences of each URL
    #create a ratio based on the size of the document and the times an URL can appear
    #if a given URL goes beyond that ratio, it's plagiarized
